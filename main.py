import yaml
from docx import *
import argparse

def getSummary(obj):
    if 'summary' in obj.keys():
        return obj['summary']
    else:
        return ''

def getTitle(obj):
    if 'title' in obj.keys():
        return obj['title']
    else:
        return ''

def getDescription(obj):
    if 'description' in obj.keys():
        return obj['description']
    else:
        return ''

def getType(obj):
    global keywords
    global obj_stack
    res = ""
    if 'type' in obj.keys():
        res = obj['type']
        if obj['type'] == 'array' and 'items' in obj.keys():
            if '$ref' in obj['items'].keys():
                res += " " + keywords['items'] + " " + str(obj['items']['$ref']) + ". "
                obj_stack.append(obj['items']['$ref'])
            else:
                if 'type' in obj['items'].keys():
                    res += " " + keywords['items'] + " " + str(obj['items']['type']) + ". "
                if 'enum' in obj['items'].keys():
                    res += keywords['enum'] + ': ' + ";".join(obj['items']['enum']) + ". "
                if 'enum' in obj.keys():
                    res += keywords['enum'] + ': ' + ";".join(obj['enum']) + ". "
                if 'default' in obj['items'].keys():
                    res += keywords['default'] + " " + str(obj['items']['default']) + ". "
                if 'collectionFormat' in obj['items'].keys():
                    res += keywords['collectionFormat'] + " " + str(obj['items']['collectionFormat']) + ". "
        if '$ref' in obj.keys():
            res += keywords['items'] + " " + str(obj['$ref']) + ". "
            obj_stack.append(obj['ref'])
        if 'minimum' in obj.keys():
            res += keywords['minimum'] + ": " + str(obj['minimum']) + '; '
        if 'maximum' in obj.keys():
            res += keywords['maximum'] + ": " + str(obj['maximum']) + '; '
        if 'minLength' in obj.keys():
            res += keywords['minLength'] + ": " + str(obj['min_length']) + '; '
        if 'maxLength' in obj.keys():
            res += keywords['maxLength'] + ": " + str(obj['max_length']) + '; '
        return res
    else:
        return res

def getFormat(obj):
    if 'format' in obj.keys():
        return obj['format']
    else:
        return ''
def getXml(obj):
    if 'xml' in obj.keys():
        return obj['xml']
    else:
        return ''

def getVersion(obj):
    if 'version' in obj.keys():
        return str(obj['version'])
    else:
        return ''

def getIn(obj):
    if 'in' in obj.keys():
        return obj['in']
    else:
        return ''

def getRequired(obj):
    if 'required' in obj.keys():
        if obj['required'] == 'true' or obj['required'] == 'True':
            return True
    return False

def isRequired(obj, param):
    if 'required' in obj.keys():
        return param in obj['required']
    return False

def getOperationId(obj):
    if 'operationId' in obj.keys():
        return obj['operationId']
    else:
        return ''

def getName(obj):
    if 'name' in obj.keys():
        return obj['name']
    else:
        return ''

def getConsumes(obj):
    if 'consumes' in obj.keys():
        return ';'.join(obj['consumes'])
    else:
        return ''

def getTags(obj):
    if 'tags' in obj.keys():
        return ';'.join(obj['tags'])
    else:
        return ''

def getProduces(obj):
    if 'produces' in obj.keys():
        return ';'.join(obj['produces'])
    else:
        return ''

def getSchema(obj):
    global obj_stack
    if 'schema' in obj.keys():
        if '$ref' in obj['schema'].keys():
            obj_stack.append(obj['schema']['$ref'])
            return obj['schema']['$ref']
        else:
            return ''
    else:
        return ''

def getObject(yaml_doc, path):
    cur_obj = yaml_doc
    for i in path.split('/'):
        if i != '#':
            cur_obj = cur_obj[i]
    return (cur_obj)

def descObject(obj, path):
    global keywords
    global obj_stack
    res_doc.add_heading(str(path), 2)
    resp_table = res_doc.add_table(rows=1, cols=7)
    hdr_cells = resp_table.rows[0].cells
    hdr_cells[0].text = keywords['param_name']
    hdr_cells[1].text = keywords['description']
    hdr_cells[2].text = keywords['summary']
    hdr_cells[3].text = keywords['type']
    hdr_cells[4].text = keywords['format']
    hdr_cells[5].text = keywords['required']
    hdr_cells[6].text = keywords['xml']
    if 'properties' in obj.keys():
        for cur_param in obj['properties']:
            row_cells = resp_table.add_row().cells
            row_cells[0].text = cur_param
            row_cells[1].text = getDescription(obj['properties'][cur_param])
            row_cells[2].text = getSummary(obj['properties'][cur_param])
            if 'ref' not in obj['properties'][cur_param].keys():
                row_cells[3].text = getType(obj['properties'][cur_param])
            else:
                row_cells[3].text = getType(obj['properties'][cur_param]) + keywords['items'] + getSchema(obj['properties'][cur_param])
            row_cells[4].text = getFormat(obj['properties'][cur_param])
            if isRequired(obj, cur_param):
                row_cells[5].text = keywords['is_true']
            else:
                row_cells[5].text = keywords['is_false']
            row_cells[6].text = getXml(obj['properties'][cur_param])

parser = argparse.ArgumentParser()
parser.add_argument('-i', help='input file', default='demo/swagger.yaml')
parser.add_argument('-o', help='output file', default='demo/swagger.docx')
parser.add_argument('-l', help='language file', default='config/en_lang.yaml')
parser.add_argument('-e', help='encoding of language file', default='utf-8')
args = parser.parse_args()

source_doc = args.i
out_doc = args.o
lang_setting = args.l
setting_encoding = args.e

stream = open(source_doc, 'r')
yaml_doc = yaml.load(stream)

res_doc = Document()
obj_stack = []
stream_lang = open(lang_setting, 'r', encoding=setting_encoding)
keywords = yaml.load(stream_lang)

if 'info' in yaml_doc.keys():
    doc_title = getTitle(yaml_doc['info'])
    if len(doc_title) > 0:
        res_doc.add_heading(doc_title, 0)
    else:
        res_doc.add_heading(source_doc, 0)
    res_doc.add_paragraph(keywords['version'] + ": " + getVersion(yaml_doc['info']))
    res_doc.add_paragraph(keywords['description'] + ": " + getDescription(yaml_doc['info']))
else:
    res_doc.add_heading(source_doc, 0)

paths = yaml_doc['paths']

for cur_method in paths:
    res_doc.add_heading(cur_method, 1)
    for cur_http_method in paths[cur_method]:
        res_doc.add_paragraph(keywords['method'] + ": " + cur_http_method)
        cur_obj = paths[cur_method][cur_http_method]
        res_doc.add_paragraph(keywords['summary'] + ": " + getSummary(cur_obj))
        res_doc.add_paragraph(keywords['tags'] + ": " + getTags(cur_obj))
        res_doc.add_paragraph(keywords['description'] + ": " + getDescription(cur_obj))
        res_doc.add_paragraph(keywords['operationId'] + ": " + getOperationId(cur_obj))
        res_doc.add_paragraph(keywords['consumes'] + ": " + getConsumes(cur_obj))
        res_doc.add_paragraph(keywords['produces'] + ": " + getProduces(cur_obj))
        if 'parameters' in cur_obj.keys():
            if len(cur_obj['parameters']) > 0:
                resp_table = res_doc.add_table(rows=1, cols=6)
                hdr_cells = resp_table.rows[0].cells
                hdr_cells[0].text = keywords['param_name']
                hdr_cells[1].text = keywords['description']
                hdr_cells[2].text = keywords['in']
                hdr_cells[3].text = keywords['type']
                hdr_cells[4].text = keywords['format']
                hdr_cells[5].text = keywords['required']
                for cur_param in cur_obj['parameters']:
                    row_cells = resp_table.add_row().cells
                    row_cells[0].text = getName(cur_param)
                    row_cells[1].text = getDescription(cur_param)
                    row_cells[2].text = getIn(cur_param)
                    row_cells[3].text = getType(cur_param)
                    row_cells[4].text = getFormat(cur_param)
                    if getRequired(cur_param):
                        row_cells[5].text = keywords['is_true']
                    else:
                        row_cells[5].text = keywords['is_false']

        res_doc.add_paragraph(keywords['responses'] + ": " )
        if 'responses' in cur_obj.keys():
            resp_table = res_doc.add_table(rows=1, cols=3)
            hdr_cells = resp_table.rows[0].cells
            hdr_cells[0].text = keywords['response_code']
            hdr_cells[1].text = keywords['description']
            hdr_cells[2].text = keywords['schema']
            if len(cur_obj['responses'])> 0:
                for cur_responce in cur_obj['responses']:
                    row_cells = resp_table.add_row().cells
                    row_cells[0].text = cur_responce
                    row_cells[1].text = getDescription(cur_obj['responses'][cur_responce])
                    row_cells[2].text = getSchema(cur_obj['responses'][cur_responce])
    printed_obj = []
    while len(obj_stack) > 0:
        i = obj_stack.pop()
        if i not in printed_obj:
            descObject( getObject(yaml_doc, i) , i)
            printed_obj.append(i)

res_doc.add_paragraph()
res_doc.save(out_doc)